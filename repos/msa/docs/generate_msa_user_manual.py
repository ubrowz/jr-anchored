"""
generate_msa_user_manual.py

Generates repos/msa/docs/msa_user_manual.docx — the engineer-facing user guide
for the JR Validated Environment MSA Module.

Run with:
    python3 repos/msa/docs/generate_msa_user_manual.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "msa_user_manual.docx")

NAVY  = "1A1A2E"
GRAY  = "F2F2F2"
WHITE = "FFFFFF"
INFO  = "E8F0F7"
WARN  = "FFF3CD"
GREEN = "E8F5E9"

# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def set_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def header_row(table, cols, widths=None):
    row = table.rows[0]
    for i, text in enumerate(cols):
        cell = row.cells[i]
        set_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)


def data_row(table, row_idx, values, shade=True, bold_first=False):
    row = table.rows[row_idx]
    fill = GRAY if (shade and row_idx % 2 == 1) else WHITE
    for i, text in enumerate(values):
        cell = row.cells[i]
        set_shading(cell, fill)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        if bold_first and i == 0:
            run.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(13)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def para(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = "Calibri"


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.name = "Calibri"


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.4)
    run = p.add_run(text)
    run.font.size = Pt(9); run.font.name = "Courier New"


def box(doc, fill, label, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_shading(cell, fill)
    p = cell.paragraphs[0]
    p.clear()
    r1 = p.add_run(label + "  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.name = "Calibri"
    r2 = p.add_run(text)
    r2.font.size = Pt(10); r2.font.name = "Calibri"
    doc.add_paragraph()


def info(doc, label, text):  box(doc, INFO, label, text)
def warn(doc, label, text):  box(doc, WARN, label, text)
def tip(doc, label, text):   box(doc, GREEN, label, text)


def script_section(doc, number, script_name, when, args_table, example_cmd, output_note):
    """Standard subsection layout for each script."""
    h2(doc, f"{number}  {script_name}")
    h3(doc, "When do I use this?")
    para(doc, when)
    h3(doc, "What do I need?")
    # args table: list of (argument, description)
    tbl = doc.add_table(rows=len(args_table) + 1, cols=2)
    tbl.style = "Table Grid"
    header_row(tbl, ["Argument", "What to provide"], [1.8, 5.1])
    for i, (arg, desc) in enumerate(args_table, start=1):
        data_row(tbl, i, [arg, desc], bold_first=True)
    doc.add_paragraph()
    h3(doc, "Example")
    code(doc, example_cmd)
    h3(doc, "What to look for in the output")
    para(doc, output_note)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()
section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1.18)
section.right_margin  = Inches(0.984)
section.top_margin    = Inches(0.984)
section.bottom_margin = Inches(0.984)

# ---------------------------------------------------------------------------
# Cover
# ---------------------------------------------------------------------------
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(40)
cover.paragraph_format.space_after  = Pt(6)
r = cover.add_run("JR Validated Environment \u2014 MSA Module")
r.bold = True; r.font.size = Pt(18); r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
r = sub.add_run("User Guide for Measurement System Analysis")
r.bold = True; r.font.size = Pt(13); r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Version 1.0  |  Date: 2026-03-18  |  Audience: Design Engineers")
r.font.size = Pt(10); r.font.name = "Calibri"
doc.add_paragraph()

# ---------------------------------------------------------------------------
# Section 1 — How to Use This Guide
# ---------------------------------------------------------------------------
h1(doc, "1. How to Use This Guide")
para(doc,
    "This guide describes the five MSA scripts available in the JR Validated Environment. "
    "Each script covers a distinct phase or approach to qualifying a measurement system. "
    "Choose the workflow that matches your test type and stage in development.")

para(doc, "The five MSA scripts cover three workflows:")
bullet(doc, "Before using a gauge for continuous measurements: Type 1 study \u2192 Gauge R&R \u2192 Linearity & Bias")
bullet(doc, "For destructive/semi-destructive testing: Nested Gauge R&R")
bullet(doc, "For attribute (pass/fail) inspection systems: Attribute Agreement Analysis")

info(doc, "INFO:",
    "Run Type 1 first (Section 2.3) before a full Gauge R&R (Section 2.1). "
    "Type 1 uses a single reference part and is quick to run. "
    "Gauge R&R requires multiple operators and is a larger study.")

# ---------------------------------------------------------------------------
# Section 2 — Quick Reference Table
# ---------------------------------------------------------------------------
h1(doc, "2. Quick Reference Table")
para(doc, "Find the row that matches what you are trying to do, then go to the section listed.")
doc.add_paragraph()

qr_rows = [
    ("Qualify a continuous gauge \u2014 full study (operators \u00d7 parts \u00d7 replicates)",
     "jrc_msa_gauge_rr", "2.1"),
    ("Qualify a gauge for destructive testing",
     "jrc_msa_nested_grr", "2.2"),
    ("Quick gauge capability check (single reference part)",
     "jrc_msa_type1", "2.3"),
    ("Assess gauge accuracy across its measurement range",
     "jrc_msa_linearity_bias", "2.4"),
    ("Qualify an attribute (pass/fail) inspection system",
     "jrc_msa_attribute", "2.5"),
]

tbl = doc.add_table(rows=len(qr_rows) + 1, cols=3)
tbl.style = "Table Grid"
header_row(tbl, ["I want to\u2026", "Script", "Section"], [3.8, 2.2, 0.9])
for i, row_data in enumerate(qr_rows, start=1):
    data_row(tbl, i, list(row_data))
doc.add_paragraph()

# ---------------------------------------------------------------------------
# Section 3 — Understanding the Results
# ---------------------------------------------------------------------------
h1(doc, "3. Understanding the Results")

h2(doc, "3.1  %GRR (Gauge R&R and Nested GRR)")
para(doc,
    "%GRR tells you what fraction of the total observed variation comes from the measurement "
    "system itself. Industry rule: < 10% Acceptable, 10\u201330% Marginal, > 30% Unacceptable. "
    "The ndc (number of distinct categories) must be \u2265 5.")

h2(doc, "3.2  Cg and Cgk (Type 1 study)")
para(doc,
    "Cg measures the gauge\u2019s spread relative to the tolerance. Cgk adjusts for bias. "
    "Both must be \u2265 1.33. If Cgk < Cg, the gauge has a systematic offset \u2014 "
    "investigate and re-centre before proceeding.")

h2(doc, "3.3  Linearity and Bias")
para(doc,
    "A flat regression line (slope \u2248 0) means the gauge bias is consistent across its range. "
    "A significant slope means the gauge reads high at one end and low at the other \u2014 "
    "it must be re-calibrated or its use restricted to a narrower range.")

h2(doc, "3.4  Kappa (Attribute Agreement)")
para(doc,
    "Kappa measures inter-rater agreement beyond chance. \u2265 0.9 Acceptable, 0.7\u20130.9 Marginal, "
    "< 0.7 Unacceptable. Fleiss\u2019 Kappa covers all appraisers simultaneously; Cohen\u2019s Kappa "
    "covers each appraiser internally (trial 1 vs trial 2) or vs reference.")

# ---------------------------------------------------------------------------
# Section 4 — Per-Script Reference
# ---------------------------------------------------------------------------
h1(doc, "4. Per-Script Reference")

# 4.1 jrc_msa_gauge_rr
script_section(doc,
    "4.1", "jrc_msa_gauge_rr \u2014 Standard Gauge R&R",
    "Use when the same physical parts can be measured by all operators. This is the standard "
    "Gauge R&R study: multiple operators each measure the same set of parts multiple times. "
    "Use this before relying on a measurement system for design verification.",
    [
        ("data.csv",
         "CSV with columns: part, operator, value. Design must be balanced: every part-operator "
         "combination must have the same number of replicates (typically 2\u20133). "
         "Minimum: 2 parts, 2 operators, 2 replicates per cell."),
        ("--tolerance <value>",
         "Optional. Process tolerance (USL \u2212 LSL). When supplied, %GRR is also expressed "
         "as a fraction of the tolerance width."),
    ],
    "jrc_msa_gauge_rr ~/msa/gauge_study.csv --tolerance 0.50",
    "Check %GRR in the Verdict section. Look for %GRR < 10% (Acceptable) or < 30% (Marginal). "
    "The ndc (number of distinct categories) must be \u2265 5. The four-panel PNG shows components "
    "of variation, by-part, by-operator, and an operator\u00d7part interaction plot \u2014 a flat "
    "interaction plot means operator differences are consistent across all parts."
)

# 4.2 jrc_msa_nested_grr
script_section(doc,
    "4.2", "jrc_msa_nested_grr \u2014 Nested Gauge R&R (Destructive)",
    "Use when each operator receives their own set of specimens that cannot be shared \u2014 "
    "for example, tensile testing, chemical analysis, or any test that destroys or permanently "
    "alters the specimen.",
    [
        ("data.csv",
         "CSV with columns: operator, part, replicate, value. Parts are nested within operators "
         "\u2014 each part belongs to exactly one operator. Part IDs may be reused across operators "
         "(operator A\u2019s part 1 and operator B\u2019s part 1 are different physical specimens). "
         "Each operator must have the same number of parts and each operator-part cell must have "
         "the same number of replicates (\u2265 2)."),
        ("--tolerance <value>",
         "Optional. Process tolerance (USL \u2212 LSL)."),
    ],
    "jrc_msa_nested_grr ~/msa/destructive_study.csv --tolerance 10.0",
    "Interpret %GRR the same way as a standard Gauge R&R (< 10% Acceptable, < 30% Marginal). "
    "Note: because parts are not shared, the part-to-part estimate is confounded with the "
    "operator batch \u2014 the analysis assumes all operators received specimens from the same population."
)
warn(doc, "WARNING:",
    "Nested GRR limitation: Because parts are not crossed across operators, part-to-part variation "
    "and reproducibility are partially aliased. The %GRR estimate is conservative. If parts can "
    "physically be shared without being destroyed, run jrc_msa_gauge_rr instead.")

# 4.3 jrc_msa_type1
script_section(doc,
    "4.3", "jrc_msa_type1 \u2014 Type 1 Gauge Study",
    "Use for a quick initial gauge capability check before committing to a full Gauge R&R study. "
    "A single reference part of known value is measured repeatedly under the same conditions. "
    "The study quantifies the gauge\u2019s spread (Cg) and any systematic bias (Cgk).",
    [
        ("data.csv",
         "CSV with a \u2018value\u2019 column. Optionally an \u2018id\u2019 column for the run chart x-axis. "
         "Minimum 10 measurements; 25\u201350 are typical."),
        ("--reference <value>",
         "Required. The known true value of the reference part "
         "(e.g. from a certified reference artefact)."),
        ("--tolerance <value>",
         "Required. The process tolerance (USL \u2212 LSL) the gauge is intended to support."),
    ],
    "jrc_msa_type1 ~/msa/type1_study.csv --reference 50.0 --tolerance 1.0",
    "Cg \u2265 1.33 and Cgk \u2265 1.33 are required. If Cg \u2265 1.33 but Cgk < 1.33, the gauge is precise "
    "but offset \u2014 check calibration. If both are below 1.33, the gauge variability is too large "
    "for this tolerance. The run chart shows whether variability is stable over time; the histogram "
    "shows whether it follows a normal distribution."
)

# 4.4 jrc_msa_linearity_bias
script_section(doc,
    "4.4", "jrc_msa_linearity_bias \u2014 Linearity and Bias",
    "Use to check whether a gauge has a consistent bias across its measurement range, or whether "
    "its accuracy changes at low vs high values. Run after Type 1 and Gauge R&R for a complete "
    "MSA picture.",
    [
        ("data.csv",
         "CSV with columns: part, reference, value. Each part must have a fixed reference value. "
         "At least 2 (ideally 5+) distinct reference levels covering the gauge\u2019s operating range. "
         "Multiple replicates per reference level are recommended."),
        ("--tolerance <value>",
         "Optional. Process tolerance."),
    ],
    "jrc_msa_linearity_bias ~/msa/linearity_study.csv --tolerance 5.0",
    "A non-significant slope (p > 0.05) means bias is consistent across the range \u2014 acceptable. "
    "A significant slope means the gauge reads differently at different reference levels \u2014 "
    "investigate calibration or restrict the gauge\u2019s working range. Per-part bias bars show "
    "which reference levels have the largest offset."
)

# 4.5 jrc_msa_attribute
h2(doc, "4.5  jrc_msa_attribute \u2014 Attribute Agreement Analysis")
h3(doc, "When do I use this?")
para(doc,
    "Use to qualify an inspection system that produces categorical results "
    "(Pass/Fail, Grade A/B/C, Accept/Reject). Multiple appraisers each inspect the same set "
    "of parts multiple times.")
h3(doc, "What do I need?")
attr_args = [
    ("data.csv",
     "CSV with columns: part, appraiser, trial, rating. The design must be balanced: every "
     "appraiser-part combination must have the same number of trials (\u2265 2). The \u2018rating\u2019 "
     "column contains the categorical result (any text or number). Optionally include a "
     "\u2018reference\u2019 column with the known correct rating for each part \u2014 this enables "
     "the vs-reference Kappa calculation."),
]
tbl = doc.add_table(rows=len(attr_args) + 1, cols=2)
tbl.style = "Table Grid"
header_row(tbl, ["Argument", "What to provide"], [1.8, 5.1])
for i, (arg, desc) in enumerate(attr_args, start=1):
    data_row(tbl, i, [arg, desc], bold_first=True)
doc.add_paragraph()
h3(doc, "Examples")
code(doc, "jrc_msa_attribute ~/msa/attribute_study.csv")
code(doc, "jrc_msa_attribute ~/msa/attribute_study_with_ref.csv")
h3(doc, "What to look for in the output")
para(doc,
    "Within-appraiser Kappa \u2265 0.9 means each appraiser is consistent with themselves. "
    "Between-appraiser Fleiss\u2019 Kappa \u2265 0.9 means the inspection system is consistent "
    "across operators. If a reference column is present, each appraiser\u2019s Kappa vs reference "
    "measures accuracy (not just consistency). Verdict: \u2265 0.9 Acceptable, 0.7\u20130.9 Marginal, "
    "< 0.7 Unacceptable.")
tip(doc, "TIP:",
    "The reference column should contain the \u2018ground truth\u2019 rating \u2014 for example, ratings "
    "from a subject matter expert or certified inspector. Without it, the analysis can only tell "
    "you whether appraisers agree with each other, not whether they are correct.")

# ---------------------------------------------------------------------------
# Section 5 — Preparing Your Data
# ---------------------------------------------------------------------------
h1(doc, "5. Preparing Your Data")
para(doc,
    "Each script expects a CSV file. The table below lists the required and optional columns "
    "for each script.")
doc.add_paragraph()

csv_rows = [
    ("jrc_msa_gauge_rr",      "part, operator, value",              "\u2014"),
    ("jrc_msa_nested_grr",    "operator, part, replicate, value",   "\u2014"),
    ("jrc_msa_type1",         "value",                              "id"),
    ("jrc_msa_linearity_bias","part, reference, value",             "\u2014"),
    ("jrc_msa_attribute",     "part, appraiser, trial, rating",     "reference"),
]

tbl = doc.add_table(rows=len(csv_rows) + 1, cols=3)
tbl.style = "Table Grid"
header_row(tbl, ["Script", "Required columns", "Optional columns"], [2.0, 3.7, 1.2])
for i, row_data in enumerate(csv_rows, start=1):
    data_row(tbl, i, list(row_data), bold_first=True)
doc.add_paragraph()

para(doc,
    "Column names are case-insensitive. Extra columns are ignored. Numeric values may use "
    "European decimal commas (they are automatically converted). The \u2018rating\u2019 column in "
    "jrc_msa_attribute may contain any text or number.")

# ---------------------------------------------------------------------------
# Section 6 — Recommended Study Workflow
# ---------------------------------------------------------------------------
h1(doc, "6. Recommended Study Workflow")
para(doc,
    "Follow this sequence when qualifying a continuous measurement system for design verification:")

p = doc.add_paragraph(style="List Number")
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run(
    "Run jrc_msa_type1 \u2014 quick gauge capability screen (Cg, Cgk). "
    "If Cg < 1.33, stop: the gauge is not capable for this tolerance.")
run.font.size = Pt(10); run.font.name = "Calibri"

p = doc.add_paragraph(style="List Number")
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run(
    "Run jrc_msa_gauge_rr (or jrc_msa_nested_grr for destructive testing) \u2014 full %GRR study. "
    "If %GRR > 30%, stop: the gauge must be improved.")
run.font.size = Pt(10); run.font.name = "Calibri"

p = doc.add_paragraph(style="List Number")
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run(
    "Optionally run jrc_msa_linearity_bias to verify accuracy across the operating range.")
run.font.size = Pt(10); run.font.name = "Calibri"

p = doc.add_paragraph(style="List Number")
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run(
    "Document all three PNG outputs in the Measurement System Qualification record in your DHF.")
run.font.size = Pt(10); run.font.name = "Calibri"

warn(doc, "WARNING:",
    "Do not proceed to design verification if %GRR > 30% or Cg/Cgk < 1.33. "
    "Measurements from an unqualified gauge cannot support FDA submission claims.")

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
